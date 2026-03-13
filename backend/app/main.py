"""
Delivery Intelligence Platform — FastAPI application with Jinja2 server-rendered UI.
"""
from __future__ import annotations

import csv
import io
import logging
import os
import re
import uuid
from datetime import datetime
from pathlib import Path

from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import HTMLResponse, RedirectResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from backend.app.config import get_settings
from backend.app.routers import ba, jira, pm, power
from backend.app.routers import batch as batch_router
from backend.app.routers import knowledge as knowledge_router
from backend.app.routers import projects as projects_router
from backend.app.routers import qa as qa_router
from backend.app.routers import requirements as requirements_router
from backend.app.services.ba_agent import BAAgent
from backend.app.services.document_store import get_document_store
from backend.app.services.jira_client import JiraError, get_jira_client
from backend.app.services.llm_client import LLMError, get_llm_client
from backend.app.services.pm_agent import PMAgent
from backend.app.services.prompt_loader import get_prompt_loader

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(name)s — %(message)s",
)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# App & static assets
# ---------------------------------------------------------------------------
settings = get_settings()

app = FastAPI(title="Delivery Intelligence Platform", version="1.0.0")

_base = Path(__file__).resolve().parent.parent.parent  # repo root

app.mount(
    "/static",
    StaticFiles(directory=str(_base / "frontend" / "static")),
    name="static",
)

templates = Jinja2Templates(directory=str(_base / "frontend" / "templates"))

# ---------------------------------------------------------------------------
# API routers
# ---------------------------------------------------------------------------
app.include_router(ba.router)
app.include_router(pm.router)
app.include_router(jira.router)
app.include_router(power.router)
app.include_router(projects_router.router)
app.include_router(batch_router.router)
app.include_router(knowledge_router.router)
app.include_router(requirements_router.router)
app.include_router(qa_router.router)

# ---------------------------------------------------------------------------
# UI helpers
# ---------------------------------------------------------------------------


def _new_sid() -> str:
    return uuid.uuid4().hex


def _get_or_create_session(request: Request) -> str:
    sid = request.cookies.get("session_id")
    if not sid:
        sid = _new_sid()
    return sid


def _set_session(response, sid: str):
    response.set_cookie("session_id", sid, httponly=True)
    return response


# ---------------------------------------------------------------------------
# UI pages
# ---------------------------------------------------------------------------


@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    sid = _get_or_create_session(request)
    store = get_document_store()
    ws = store.load_workspace(sid)
    response = templates.TemplateResponse(
        "index.html",
        {"request": request, "session_id": sid, "workspace": ws},
    )
    response.set_cookie("session_id", sid, httponly=True)
    return response


@app.get("/ba/requirements", response_class=HTMLResponse)
async def ba_requirements_page(request: Request):
    sid = _get_or_create_session(request)
    store = get_document_store()
    ws = store.load_workspace(sid)
    managed_projects = store.load_project_registry()
    response = templates.TemplateResponse(
        "ba_requirements.html",
        {
            "request": request,
            "session_id": sid,
            "workspace": ws,
            "managed_projects": managed_projects,
            "result": None,
            "error": None,
        },
    )
    response.set_cookie("session_id", sid, httponly=True)
    return response


@app.post("/ba/requirements", response_class=HTMLResponse)
async def ba_requirements_post(
    request: Request,
    action: str = Form(...),
    session_id: str = Form(...),
    jira_project_key: str = Form(""),
    ba_source: str = Form("new"),
    jira_story_key: str = Form(""),
    raw_notes: str = Form(""),
    edit_instruction: str = Form(""),
):
    store = get_document_store()
    jira_client = get_jira_client()
    agent = BAAgent(
        llm=get_llm_client(),
        prompt_loader=get_prompt_loader(),
        store=store,
        jira=jira_client if jira_client.is_configured() else None,
    )
    result = None
    pulled_story = None
    error = None

    # Persist mode/project selection; handle state clearing on mode switch
    ws = store.load_workspace(session_id)
    ws.jira_project_key = jira_project_key.strip()
    prev_source = ws.ba_source
    ws.ba_source = ba_source if ba_source in ("new", "existing_story") else "new"
    if ws.ba_source == "new" and prev_source == "existing_story":
        ws.jira_story_key = ""
        ws.pulled_jira_story = None
        ws.readiness_report = None
    elif ws.ba_source == "existing_story":
        new_key = jira_story_key.strip()
        if new_key and new_key != ws.jira_story_key:
            ws.jira_story_key = new_key
            ws.pulled_jira_story = None
            ws.readiness_report = None
    store.save_workspace(ws)

    try:
        if action == "generate" and ws.ba_source == "new":
            result = agent.generate_requirements(session_id, raw_notes)
        elif action == "update" and ws.ba_source == "new":
            result = agent.update_requirements(session_id, edit_instruction)
        elif action == "pull_story":
            key = ws.jira_story_key or jira_story_key.strip()
            if not key:
                error = "Enter a Jira story key (e.g. PROJ-123) before pulling."
            else:
                pulled_story = agent.pull_jira_story(session_id, key)
    except (LLMError, ValueError, FileNotFoundError, JiraError) as exc:
        error = str(exc)

    ws = store.load_workspace(session_id)
    managed_projects = store.load_project_registry()
    response = templates.TemplateResponse(
        "ba_requirements.html",
        {
            "request": request,
            "session_id": session_id,
            "workspace": ws,
            "managed_projects": managed_projects,
            "result": result,
            "pulled_story": pulled_story,
            "error": error,
            "submitted_raw_notes": raw_notes,
        },
    )
    response.set_cookie("session_id", session_id, httponly=True)
    return response


@app.get("/ba/stories", response_class=HTMLResponse)
async def ba_stories_page(request: Request):
    sid = _get_or_create_session(request)
    store = get_document_store()
    ws = store.load_workspace(sid)
    response = templates.TemplateResponse(
        "story_set.html",
        {"request": request, "session_id": sid, "workspace": ws, "result": None, "error": None},
    )
    response.set_cookie("session_id", sid, httponly=True)
    return response


@app.post("/ba/stories", response_class=HTMLResponse)
async def ba_stories_post(
    request: Request,
    action: str = Form(...),
    session_id: str = Form(...),
    edit_instruction: str = Form(""),
):
    store = get_document_store()
    agent = BAAgent(llm=get_llm_client(), prompt_loader=get_prompt_loader(), store=store)
    result = None
    error = None
    try:
        if action == "generate":
            result = agent.generate_stories(session_id)
        elif action == "update":
            result = agent.update_stories(session_id, edit_instruction)
    except (LLMError, ValueError, FileNotFoundError) as exc:
        error = str(exc)

    ws = store.load_workspace(session_id)
    response = templates.TemplateResponse(
        "story_set.html",
        {
            "request": request,
            "session_id": session_id,
            "workspace": ws,
            "result": result,
            "error": error,
        },
    )
    response.set_cookie("session_id", session_id, httponly=True)
    return response


@app.get("/ba/readiness", response_class=HTMLResponse)
async def ba_readiness_page(request: Request):
    sid = _get_or_create_session(request)
    store = get_document_store()
    ws = store.load_workspace(sid)
    response = templates.TemplateResponse(
        "readiness.html",
        {
            "request": request,
            "session_id": sid,
            "workspace": ws,
            "result": None,
            "jira_result": None,
            "approve_result": None,
            "jira_configured": get_jira_client().is_configured(),
            "error": None,
        },
    )
    response.set_cookie("session_id", sid, httponly=True)
    return response


@app.post("/ba/readiness", response_class=HTMLResponse)
async def ba_readiness_post(
    request: Request,
    action: str = Form(...),
    session_id: str = Form(...),
    dry_run: str = Form("true"),
):
    store = get_document_store()
    jira_client = get_jira_client()
    agent = BAAgent(
        llm=get_llm_client(),
        prompt_loader=get_prompt_loader(),
        store=store,
        jira=jira_client if jira_client.is_configured() else None,
    )
    result = None
    jira_result = None
    approve_result = None
    error = None
    ws = store.load_workspace(session_id)
    try:
        if action == "check":
            if ws.ba_source == "existing_story":
                # Re-pull from Jira then run check
                if ws.jira_story_key:
                    agent.pull_jira_story(session_id, ws.jira_story_key)
                result = agent.check_readiness_from_jira_story(session_id)
            else:
                result = agent.check_readiness(session_id)
        elif action == "approve":
            approve_result = agent.approve_jira_story(session_id)
        elif action in ("jira_dry", "jira_create"):
            from backend.app.schemas import JiraCreateRequest
            from backend.app.routers.jira import create_story_set
            is_dry = action == "jira_dry"
            req = JiraCreateRequest(session_id=session_id, dry_run=is_dry)
            jira_result = await create_story_set(req, store=store, jira=jira_client)
    except (LLMError, ValueError, FileNotFoundError, JiraError) as exc:
        error = str(exc)
    except HTTPException as exc:
        error = exc.detail

    ws = store.load_workspace(session_id)
    response = templates.TemplateResponse(
        "readiness.html",
        {
            "request": request,
            "session_id": session_id,
            "workspace": ws,
            "result": result,
            "jira_result": jira_result,
            "approve_result": approve_result,
            "jira_configured": jira_client.is_configured(),
            "error": error,
        },
    )
    response.set_cookie("session_id", session_id, httponly=True)
    return response


# ---------------------------------------------------------------------------
# Requirements document upload
# ---------------------------------------------------------------------------

_ALLOWED_CONTEXT_EXTS = {".docx", ".xlsx", ".txt", ".md"}


def _extract_text(filename: str, content: bytes) -> str:
    """Extract plain text from an uploaded file."""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in _ALLOWED_CONTEXT_EXTS:
        raise ValueError(f"Unsupported file type '{ext}'. Allowed: {', '.join(sorted(_ALLOWED_CONTEXT_EXTS))}")
    if ext == ".docx":
        import io
        from docx import Document  # type: ignore[import-untyped]
        doc = Document(io.BytesIO(content))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)
    if ext == ".xlsx":
        import io
        import openpyxl  # type: ignore[import-untyped]
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join(str(v) if v is not None else "" for v in row)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)
    # .txt / .md
    return content.decode("utf-8", errors="replace")


@app.post("/ba/requirements/upload", response_class=HTMLResponse)
async def ba_requirements_upload(
    request: Request,
    session_id: str = Form(...),
    file: UploadFile = File(...),
):
    store = get_document_store()
    upload_error = None
    upload_message = None
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in _ALLOWED_CONTEXT_EXTS:
        upload_error = f"Unsupported file type. Allowed: {', '.join(sorted(_ALLOWED_CONTEXT_EXTS))}"
    else:
        try:
            content = await file.read()
            filename = os.path.basename(file.filename or "upload")
            text = _extract_text(filename, content)
            ws = store.load_workspace(session_id)
            ws.context_docs[filename] = text
            store.save_workspace(ws)
            upload_message = f"'{filename}' uploaded and will be included when generating requirements."
        except Exception as exc:
            upload_error = f"Failed to process file: {exc}"

    ws = store.load_workspace(session_id)
    response = templates.TemplateResponse(
        "ba_requirements.html",
        {
            "request": request,
            "session_id": session_id,
            "workspace": ws,
            "result": None,
            "error": None,
            "upload_message": upload_message,
            "upload_error": upload_error,
        },
    )
    response.set_cookie("session_id", session_id, httponly=True)
    return response


@app.post("/ba/requirements/remove-doc", response_class=HTMLResponse)
async def ba_requirements_remove_doc(
    request: Request,
    session_id: str = Form(...),
    filename: str = Form(...),
):
    store = get_document_store()
    ws = store.load_workspace(session_id)
    ws.context_docs.pop(filename, None)
    store.save_workspace(ws)
    response = templates.TemplateResponse(
        "ba_requirements.html",
        {
            "request": request,
            "session_id": session_id,
            "workspace": ws,
            "result": None,
            "error": None,
        },
    )
    response.set_cookie("session_id", session_id, httponly=True)
    return response


@app.get("/pm", response_class=HTMLResponse)
async def pm_page(request: Request):
    sid = _get_or_create_session(request)
    store = get_document_store()
    ws = store.load_workspace(sid)
    response = templates.TemplateResponse(
        "pm_mode.html",
        {
            "request": request,
            "session_id": sid,
            "workspace": ws,
            "result": None,
            "error": None,
            "submitted_query": "",
        },
    )
    response.set_cookie("session_id", sid, httponly=True)
    return response


@app.post("/pm", response_class=HTMLResponse)
async def pm_post(
    request: Request,
    session_id: str = Form(...),
    query: str = Form(...),
):
    pm_agent = PMAgent(
        llm=get_llm_client(),
        prompt_loader=get_prompt_loader(),
        jira=get_jira_client(),
    )
    result = None
    error = None
    try:
        result = pm_agent.query(session_id, query)
    except JiraError as exc:
        error = str(exc)
    except (LLMError, ValueError, FileNotFoundError) as exc:
        error = str(exc)
    except Exception as exc:
        error = f"Unexpected error: {exc}"

    store = get_document_store()
    ws = store.load_workspace(session_id)
    response = templates.TemplateResponse(
        "pm_mode.html",
        {
            "request": request,
            "session_id": session_id,
            "workspace": ws,
            "result": result,
            "error": error,
            "submitted_query": query,
        },
    )
    response.set_cookie("session_id", session_id, httponly=True)
    return response


@app.post("/pm/export-csv")
async def pm_export_csv(
    session_id: str = Form(...),
    jql: str = Form(...),
):
    del session_id  # Reserved for future session-scoped export audit/history.

    jql_value = jql.strip()
    if not jql_value:
        raise HTTPException(status_code=400, detail="JQL is required for CSV export.")

    jira_client = get_jira_client()
    if not jira_client.is_configured():
        raise HTTPException(
            status_code=503,
            detail="Jira is not configured. Set Jira credentials in .env before exporting CSV.",
        )

    try:
        data = jira_client.search_issues(jql_value, max_results=1000)
    except JiraError as exc:
        raise HTTPException(status_code=exc.status_code or 502, detail=str(exc)) from exc

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Key", "Type", "Summary", "Status", "Priority", "Assignee", "Jira URL"])

    for issue in data.get("issues", []):
        fields = issue.get("fields", {})
        key = issue.get("key", "")
        issue_url = f"{jira_client._base}/browse/{key}" if key else ""
        writer.writerow(
            [
                key,
                fields.get("issuetype", {}).get("name", ""),
                fields.get("summary", ""),
                fields.get("status", {}).get("name", ""),
                (fields.get("priority") or {}).get("name", ""),
                (fields.get("assignee") or {}).get("displayName", ""),
                issue_url,
            ]
        )

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"jira_export_{timestamp}.csv"
    return Response(
        content=output.getvalue(),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/power", response_class=HTMLResponse)
async def power_page(request: Request):
    sid = _get_or_create_session(request)
    store = get_document_store()
    ws = store.load_workspace(sid)
    response = templates.TemplateResponse(
        "power_mode.html",
        {"request": request, "session_id": sid, "workspace": ws},
    )
    response.set_cookie("session_id", sid, httponly=True)
    return response


# ---------------------------------------------------------------------------
# Global exception handler
# ---------------------------------------------------------------------------


@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    logger.exception("Unhandled exception: %s", exc)
    return templates.TemplateResponse(
        "error.html",
        {"request": request, "detail": str(exc)},
        status_code=500,
    )


# ---------------------------------------------------------------------------
# Manage Projects UI
# ---------------------------------------------------------------------------


def _manage_projects_context(request: Request, sid: str, target: str, msg: str = "", error: str = ""):
    store = get_document_store()
    projects = store.load_project_registry()
    # Determine resolved target key (None = default)
    key = None if target == "default" else target if target else None
    checklist = store.get_checklist_content_response(key=key)
    history = store.list_checklist_versions(key=key)
    field_config = store.load_field_config(target) if (target and target != "default") else []
    return {
        "request": request,
        "session_id": sid,
        "projects": projects,
        "target": target or "default",
        "checklist": checklist,
        "history": history,
        "field_config": field_config,
        "msg": msg,
        "error": error,
    }


@app.get("/manage-projects", response_class=HTMLResponse)
async def manage_projects_page(request: Request, target: str = "default"):
    sid = _get_or_create_session(request)
    msg = request.query_params.get("msg", "")
    ctx = _manage_projects_context(request, sid, target, msg=msg)
    resp = templates.TemplateResponse("manage_projects.html", ctx)
    resp.set_cookie("session_id", sid, httponly=True)
    return resp


@app.post("/manage-projects/add", response_class=HTMLResponse)
async def manage_projects_add(request: Request, key: str = Form(...)):
    sid = _get_or_create_session(request)
    key = key.strip().upper()
    try:
        from backend.app.services.document_store import validate_project_key
        validate_project_key(key)
    except ValueError as exc:
        ctx = _manage_projects_context(request, sid, "default", error=str(exc))
        resp = templates.TemplateResponse("manage_projects.html", ctx)
        resp.set_cookie("session_id", sid, httponly=True)
        return resp
    # Delegate to API
    store = get_document_store()
    jira_client = get_jira_client()
    project_name = ""
    if jira_client.is_configured():
        try:
            info = jira_client.get_project(key)
            project_name = info.get("name", "")
        except Exception:
            pass
    _, already_existed = store.add_project(key, name=project_name)
    msg = "already_registered" if already_existed else "added"
    return RedirectResponse(url=f"/manage-projects?target={key}&msg={msg}", status_code=303)


@app.post("/manage-projects/remove", response_class=HTMLResponse)
async def manage_projects_remove(request: Request, key: str = Form(...)):
    sid = _get_or_create_session(request)
    store = get_document_store()
    store.remove_project(key)
    return RedirectResponse(url="/manage-projects?target=default&msg=removed", status_code=303)


@app.post("/manage-projects/checklist/save", response_class=HTMLResponse)
async def manage_projects_checklist_save(
    request: Request,
    target: str = Form(...),
    content: str = Form(...),
):
    sid = _get_or_create_session(request)
    store = get_document_store()
    key = None if target == "default" else target
    try:
        store.save_checklist(key=key, content=content)
    except ValueError as exc:
        ctx = _manage_projects_context(request, sid, target, error=str(exc))
        resp = templates.TemplateResponse("manage_projects.html", ctx)
        resp.set_cookie("session_id", sid, httponly=True)
        return resp
    return RedirectResponse(url=f"/manage-projects?target={target}&msg=saved", status_code=303)


@app.post("/manage-projects/checklist/restore", response_class=HTMLResponse)
async def manage_projects_checklist_restore(
    request: Request,
    target: str = Form(...),
    version: int = Form(...),
):
    sid = _get_or_create_session(request)
    store = get_document_store()
    key = None if target == "default" else target
    try:
        store.restore_checklist_version(key=key, version=version)
    except (FileNotFoundError, ValueError) as exc:
        ctx = _manage_projects_context(request, sid, target, error=str(exc))
        resp = templates.TemplateResponse("manage_projects.html", ctx)
        resp.set_cookie("session_id", sid, httponly=True)
        return resp
    return RedirectResponse(url=f"/manage-projects?target={target}&msg=restored", status_code=303)


@app.post("/manage-projects/checklist/delete-files", response_class=HTMLResponse)
async def manage_projects_checklist_delete(request: Request, target: str = Form(...)):
    sid = _get_or_create_session(request)
    if target == "default":
        ctx = _manage_projects_context(request, sid, target, error="The default checklist files cannot be deleted.")
        resp = templates.TemplateResponse("manage_projects.html", ctx)
        resp.set_cookie("session_id", sid, httponly=True)
        return resp
    store = get_document_store()
    try:
        store.delete_checklist_files(target)
    except ValueError as exc:
        ctx = _manage_projects_context(request, sid, target, error=str(exc))
        resp = templates.TemplateResponse("manage_projects.html", ctx)
        resp.set_cookie("session_id", sid, httponly=True)
        return resp
    return RedirectResponse(url=f"/manage-projects?target=default&msg=deleted", status_code=303)


_MAX_FIELD_MAPPINGS = 20


@app.post("/manage-projects/fields/add", response_class=HTMLResponse)
async def manage_projects_fields_add(
    request: Request,
    key: str = Form(...),
    field_id: str = Form(...),
    label: str = Form(...),
):
    from backend.app.schemas import FieldMapping
    from backend.app.services.document_store import validate_project_key
    sid = _get_or_create_session(request)
    key = key.strip().upper()
    field_id = field_id.strip()
    label = label.strip()
    try:
        validate_project_key(key)
        if not field_id:
            raise ValueError("Field ID is required.")
        if not re.match(r"^[A-Za-z][A-Za-z0-9_]*$", field_id):
            raise ValueError(
                "Field ID must be a safe token (letters, digits, underscores; no spaces or slashes). "
                "Example: customfield_10014"
            )
        if not label:
            raise ValueError("Label is required.")
    except ValueError as exc:
        ctx = _manage_projects_context(request, sid, key, error=str(exc))
        resp = templates.TemplateResponse("manage_projects.html", ctx)
        resp.set_cookie("session_id", sid, httponly=True)
        return resp
    store = get_document_store()
    mappings = store.load_field_config(key)
    # Upsert by field_id (allows label rename); reject duplicate labels (case/whitespace normalized)
    label_norm = " ".join(label.split())
    existing_labels = {" ".join(m.label.split()).lower() for m in mappings if m.field_id != field_id}
    if label_norm.lower() in existing_labels:
        ctx = _manage_projects_context(
            request, sid, key, error=f"Label '{label_norm}' is already used by another field mapping."
        )
        resp = templates.TemplateResponse("manage_projects.html", ctx)
        resp.set_cookie("session_id", sid, httponly=True)
        return resp
    mappings = [m for m in mappings if m.field_id != field_id]
    if len(mappings) >= _MAX_FIELD_MAPPINGS:
        ctx = _manage_projects_context(
            request, sid, key, error=f"Maximum of {_MAX_FIELD_MAPPINGS} custom field mappings per project."
        )
        resp = templates.TemplateResponse("manage_projects.html", ctx)
        resp.set_cookie("session_id", sid, httponly=True)
        return resp
    mappings.append(FieldMapping(field_id=field_id, label=label_norm))
    store.save_field_config(key, mappings)
    return RedirectResponse(url=f"/manage-projects?target={key}&msg=field_added", status_code=303)


@app.post("/manage-projects/fields/remove", response_class=HTMLResponse)
async def manage_projects_fields_remove(
    request: Request,
    key: str = Form(...),
    field_id: str = Form(...),
):
    sid = _get_or_create_session(request)
    key = key.strip().upper()
    store = get_document_store()
    mappings = store.load_field_config(key)
    mappings = [m for m in mappings if m.field_id != field_id]
    store.save_field_config(key, mappings)
    return RedirectResponse(url=f"/manage-projects?target={key}&msg=field_removed", status_code=303)

