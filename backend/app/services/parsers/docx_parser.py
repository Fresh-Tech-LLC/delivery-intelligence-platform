"""Microsoft Word (.docx) parser."""
from __future__ import annotations

from pathlib import Path
from typing import Any


def extract_text(path: Path) -> tuple[str, dict[str, Any]]:
    """Return (extracted_text, metadata_dict) for .docx files.

    Extracts paragraph text and table cell text. Requires python-docx.
    """
    from docx import Document  # type: ignore[import-untyped]

    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts), {
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "parser": "docx_parser",
    }
